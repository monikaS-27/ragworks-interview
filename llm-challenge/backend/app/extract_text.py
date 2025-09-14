# import os
# from typing import List
# from PyPDF2 import PdfReader
# from docx import Document

# def extract_text(file_path: str) -> List[str]:
#     """Extract text from PDF or DOCX and split into chunks."""
#     text = ""
#     ext = os.path.splitext(file_path)[1].lower()
    
#     if ext == ".pdf":
#         reader = PdfReader(file_path)
#         for page in reader.pages:
#             text += page.extract_text() + "\n"
#     elif ext == ".docx":
#         doc = Document(file_path)
#         for para in doc.paragraphs:
#             text += para.text + "\n"
#     else:
#         raise ValueError("Unsupported file type")

#     # Split into chunks of ~500 characters
#     chunks = [text[i:i+500] for i in range(0, len(text), 500)]
#     return chunks

# app/extract_text.py
import os
from typing import List

try:
    import docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

def extract_text(file_path: str) -> List[str]:
    """
    Extract text from .txt, .pdf, .docx files.
    Returns a list of chunks (for embedding).
    """
    if not os.path.exists(file_path):
        raise ValueError(f"File {file_path} does not exist")

    ext = file_path.split(".")[-1].lower()
    text = ""

    if ext == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    elif ext == "pdf":
        if not PyPDF2:
            raise ValueError("PyPDF2 not installed")
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
    elif ext == "docx":
        if not docx:
            raise ValueError("python-docx not installed")
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    # Split into chunks of ~500 characters for embeddings
    chunk_size = 500
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    return chunks
